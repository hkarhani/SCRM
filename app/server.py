from __future__ import annotations

import base64
import binascii
import csv
import io
import ipaddress
import json
import os
import re
import socket
import ssl
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


APP_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = Path(os.environ.get("SCRM_DATA_DIR", str(APP_ROOT / "SCR"))).resolve()
UPLOAD_DIR = DATA_DIR / "uploads"
SNAPSHOT_DIR = DATA_DIR / "snapshots"
DOCUMENTS_DIR = DATA_DIR / "Documents"
STATIC_DIR = Path(__file__).resolve().parent / "static"
DEFAULT_PROJECT_NAME = "Segment Conflict Workspace"

CONFIG_PATH = DATA_DIR / "config.json"
POLICIES_PATH = UPLOAD_DIR / "policies.xml"
SEGMENTS_PATH = UPLOAD_DIR / "segments.xml"
HOSTS_PATH = SNAPSHOT_DIR / "hosts.json"
ADMIN_SEGMENTS_PATH = SNAPSHOT_DIR / "admin_segments.json"
INSTRUCTIONS_PATH = SNAPSHOT_DIR / "manual_instructions.json"

for directory in (DATA_DIR, UPLOAD_DIR, SNAPSHOT_DIR, DOCUMENTS_DIR):
    directory.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Segment Conflict Resolution Management", version="0.1.0")


class ApiConfigPayload(BaseModel):
    base_url: str
    username: str
    password: str = ""
    verify_tls: bool = False


class RangeUpdate(BaseModel):
    segment_key: str = ""
    name: str = ""
    path: str = ""
    source_range: str = ""
    remove_range: str
    reason: str = ""
    stage: str = ""


class VisualizationPayload(BaseModel):
    range: str = ""
    title: str = ""
    png_data_url: str = ""


class ApplyRangesPayload(BaseModel):
    updates: list[RangeUpdate]
    scope: str = ""
    visualizations: list[VisualizationPayload] = []
    requested_mode: str = "read_only"


class ProtectionPayload(BaseModel):
    live_edit_enabled: bool = False


class ProjectPayload(BaseModel):
    name: str


@dataclass
class ConnectionConfig:
    base_url: str
    username: str
    password: str
    verify_tls: bool = False
    timeout_seconds: int = 30

    def normalized_base_url(self) -> str:
        value = (self.base_url or "").strip().rstrip("/")
        if not value:
            raise ValueError("Base URL is required.")
        if not value.startswith(("http://", "https://")):
            value = f"https://{value}"
        return value


class PlatformApiError(RuntimeError):
    pass


class WebApiClient:
    def __init__(self, config: ConnectionConfig):
        self.config = config
        self.base_url = config.normalized_base_url()
        self._token = ""
        self._ssl_context = None if config.verify_tls else ssl._create_unverified_context()

    def authenticate(self) -> str:
        if self._token:
            return self._token
        body = urllib.parse.urlencode({"username": self.config.username, "password": self.config.password}).encode()
        request = urllib.request.Request(
            f"{self.base_url}/api/login",
            data=body,
            headers={"Content-Type": "application/x-www-form-urlencoded"},
            method="POST",
        )
        self._token = self._open(request, "Web API login").decode("utf-8", errors="replace").strip()
        if not self._token:
            raise PlatformApiError("Web API login succeeded but returned an empty token.")
        return self._token

    def test(self) -> dict[str, Any]:
        return {"ok": bool(self.authenticate()), "base_url": self.base_url, "username": self.config.username}

    def get_hosts(self) -> dict[str, Any]:
        return self._json("GET", "/api/hosts")

    def _json(self, method: str, path: str) -> dict[str, Any]:
        token = self.authenticate()
        request = urllib.request.Request(
            f"{self.base_url}{path}",
            headers={"Authorization": token, "Accept": "application/hal+json, application/json"},
            method=method,
        )
        raw = self._open(request, path)
        try:
            return json.loads(raw.decode("utf-8"))
        except json.JSONDecodeError as exc:
            raise PlatformApiError(f"Web API returned non-JSON content for {path}.") from exc

    def _open(self, request: urllib.request.Request, label: str) -> bytes:
        return _open_url(request, label, "Web API", self.config.timeout_seconds, self._ssl_context)


class AdminApiClient:
    def __init__(self, config: ConnectionConfig):
        self.config = config
        self.base_url = config.normalized_base_url()
        self._token = ""
        self._ssl_context = None if config.verify_tls else ssl._create_unverified_context()

    def authenticate(self) -> str:
        if self._token:
            return self._token
        body = urllib.parse.urlencode(
            {
                "username": self.config.username,
                "password": self.config.password,
                "grant_type": "password",
                "client_id": "fs-oauth-client",
            }
        ).encode()
        request = urllib.request.Request(
            f"{self.base_url}/fsum/oauth2.0/token",
            data=body,
            headers={"Content-Type": "application/x-www-form-urlencoded", "Accept": "application/json"},
            method="POST",
        )
        raw = self._open(request, "Admin API OAuth token")
        token = _token_from_response(raw)
        if not token:
            raise PlatformApiError("Admin API OAuth returned an empty token.")
        self._token = token
        return token

    def test(self) -> dict[str, Any]:
        payload = self.get_segments()
        return {
            "ok": True,
            "base_url": self.base_url,
            "username": self.config.username,
            "segments": len(normalize_admin_segments(payload)),
        }

    def get_segments(self) -> dict[str, Any] | list[Any]:
        return self._json("GET", "/adminapi/segments")

    def put_segment_subtree(self, node: dict[str, Any]) -> None:
        self._json("PUT", "/adminapi/segments?forceChanges=true", payload=node, allow_empty=True)

    def apply_range_updates(self, updates: list[RangeUpdate]) -> dict[str, Any]:
        if not updates:
            raise ValueError("No range updates were provided.")
        payload = self.get_segments()
        grouped: dict[str, list[RangeUpdate]] = {}
        for update in updates:
            target_key = update.segment_key or f"{update.name}::{_normalize_path(update.path)}"
            grouped.setdefault(target_key, []).append(update)

        applied: list[dict[str, Any]] = []
        for target_key, group_updates in grouped.items():
            target = find_admin_node(payload, group_updates[0].segment_key, group_updates[0].path, group_updates[0].name)
            if not target:
                raise PlatformApiError(f"Unable to find live segment target: {target_key}")
            node = target["node"]
            ranges_key = "ranges" if isinstance(node.get("ranges"), list) or "Ranges" not in node else "Ranges"
            old_ranges = [str(item).strip() for item in node.get(ranges_key, []) if str(item).strip()] if isinstance(node.get(ranges_key), list) else []
            final_ranges = subtract_updates_from_ranges(old_ranges, group_updates)
            node[ranges_key] = final_ranges
            self.put_segment_subtree(node)
            applied.append(
                {
                    "segment_key": target["key"],
                    "name": target["name"],
                    "path": target["path"],
                    "old_ranges": old_ranges,
                    "ranges": final_ranges,
                    "removed": [item for item in old_ranges if item not in final_ranges],
                    "update_count": len(group_updates),
                }
            )
        return {"applied": applied, "segments": payload}

    def _json(
        self,
        method: str,
        path: str,
        payload: dict[str, Any] | list[Any] | None = None,
        allow_empty: bool = False,
    ) -> dict[str, Any] | list[Any]:
        token = self.authenticate()
        data = json.dumps(payload).encode("utf-8") if payload is not None else None
        request = urllib.request.Request(
            f"{self.base_url}{path}",
            data=data,
            headers={
                "Authorization": f"Bearer {token}",
                "Accept": "application/json",
                "Content-Type": "application/json",
            },
            method=method,
        )
        raw = self._open(request, path)
        if allow_empty and not raw.strip():
            return {}
        try:
            return json.loads(raw.decode("utf-8"))
        except json.JSONDecodeError as exc:
            if allow_empty:
                return {}
            raise PlatformApiError(f"Admin API returned non-JSON content for {path}.") from exc

    def _open(self, request: urllib.request.Request, label: str) -> bytes:
        return _open_url(request, label, "Admin API", self.config.timeout_seconds, self._ssl_context)


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/api/health")
async def health() -> dict[str, Any]:
    return {
        "ok": True,
        "app": "Segment Conflict Resolution Management",
        "version": app.version,
        "workspace": str(DATA_DIR),
    }


@app.get("/api/state")
async def get_state() -> dict[str, Any]:
    config = read_config()
    project_name = current_project_name(config)
    return {
        "workspace": {
            "root": str(DATA_DIR),
            "documents": str(DOCUMENTS_DIR),
        },
        "project": {
            "name": project_name,
            "slug": project_slug(project_name),
        },
        "artifacts": {
            "policies": file_meta(POLICIES_PATH),
            "segments": file_meta(SEGMENTS_PATH),
            "hosts": file_meta(HOSTS_PATH),
            "admin_segments": file_meta(ADMIN_SEGMENTS_PATH),
        },
        "config": {
            "web": masked_config(config.get("web", {})),
            "admin": masked_config(config.get("admin", {})),
        },
        "protection": {
            "live_edit_enabled": bool(config.get("live_edit_enabled", False)),
            "mode": "live_edit" if config.get("live_edit_enabled", False) else "read_only",
        },
        "instructions": read_json(INSTRUCTIONS_PATH, {}),
    }


@app.post("/api/project")
async def save_project(payload: ProjectPayload) -> dict[str, Any]:
    name = normalize_project_name(payload.name)
    config = read_config()
    config["project_name"] = name
    config["project_updated_at"] = utc_now()
    write_config(config)
    return {"ok": True, "project": {"name": name, "slug": project_slug(name)}}


@app.post("/api/upload/{artifact_type}")
async def upload_artifact(artifact_type: str, file: UploadFile = File(...)) -> dict[str, Any]:
    if artifact_type not in {"policies", "segments", "hosts"}:
        raise HTTPException(status_code=400, detail="Artifact type must be policies, segments, or hosts.")
    content = await file.read()
    if not content.strip():
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    if artifact_type == "policies":
        parsed = parse_policies_xml(content)
        path = POLICIES_PATH
        path.write_bytes(content)
        (UPLOAD_DIR / f"{artifact_type}.json").write_text(json.dumps(parsed, indent=2), encoding="utf-8")
    elif artifact_type == "segments":
        parsed = parse_segments_xml(content)
        path = SEGMENTS_PATH
        path.write_bytes(content)
        (UPLOAD_DIR / f"{artifact_type}.json").write_text(json.dumps(parsed, indent=2), encoding="utf-8")
    else:
        rows = parse_host_ip_file(content)
        parsed = {"summary": {"hosts": len(rows)}, "hosts": rows}
        HOSTS_PATH.write_text(json.dumps({"collected_at": utc_now(), "count": len(rows), "hosts": rows}, indent=2), encoding="utf-8")
    return {"ok": True, "artifact": artifact_type, "filename": file.filename, "summary": parsed.get("summary", {})}


@app.post("/api/config/{kind}")
async def save_config(kind: str, payload: ApiConfigPayload) -> dict[str, Any]:
    if kind not in {"web", "admin"}:
        raise HTTPException(status_code=400, detail="Config kind must be web or admin.")
    config = read_config()
    previous = config.get(kind, {})
    password = payload.password or reveal_password(previous.get("password", ""))
    config[kind] = {
        "base_url": payload.base_url.strip(),
        "username": payload.username.strip(),
        "password": hide_password(password),
        "verify_tls": bool(payload.verify_tls),
        "updated_at": utc_now(),
    }
    write_config(config)
    return {"ok": True, "config": masked_config(config[kind])}


@app.post("/api/protection")
async def save_protection(payload: ProtectionPayload) -> dict[str, Any]:
    config = read_config()
    config["live_edit_enabled"] = bool(payload.live_edit_enabled)
    config["protection_updated_at"] = utc_now()
    write_config(config)
    return {
        "ok": True,
        "live_edit_enabled": bool(config["live_edit_enabled"]),
        "mode": "live_edit" if config["live_edit_enabled"] else "read_only",
    }


@app.post("/api/test/{kind}")
async def test_config(kind: str, payload: ApiConfigPayload) -> dict[str, Any]:
    config = connection_from_payload_or_saved(kind, payload)
    try:
        result = WebApiClient(config).test() if kind == "web" else AdminApiClient(config).test()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    return result


@app.post("/api/collect/hosts")
async def collect_hosts(payload: ApiConfigPayload | None = None) -> dict[str, Any]:
    config = connection_from_payload_or_saved("web", payload)
    try:
        hosts_payload = WebApiClient(config).get_hosts()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    rows = normalize_hosts(hosts_payload)
    snapshot = {"collected_at": utc_now(), "count": len(rows), "hosts": rows}
    HOSTS_PATH.write_text(json.dumps(snapshot, indent=2), encoding="utf-8")
    return {"ok": True, "count": len(rows), "collected_at": snapshot["collected_at"]}


@app.post("/api/collect/admin-segments")
async def collect_admin_segments(payload: ApiConfigPayload | None = None) -> dict[str, Any]:
    config = connection_from_payload_or_saved("admin", payload)
    try:
        raw = AdminApiClient(config).get_segments()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    rows = normalize_admin_segments(raw)
    snapshot = {"collected_at": utc_now(), "count": len(rows), "raw": raw, "segments": rows}
    ADMIN_SEGMENTS_PATH.write_text(json.dumps(snapshot, indent=2), encoding="utf-8")
    return {"ok": True, "count": len(rows), "collected_at": snapshot["collected_at"]}


@app.get("/api/analysis")
async def get_analysis() -> dict[str, Any]:
    return build_analysis()


@app.post("/api/clear-data")
async def clear_loaded_data() -> dict[str, Any]:
    removed = clear_workspace_files(include_documents=True)
    if CONFIG_PATH.exists():
        CONFIG_PATH.unlink()
        removed.append(str(CONFIG_PATH.relative_to(DATA_DIR)))
    return {"ok": True, "removed": removed}


@app.get("/api/export/workspace.zip")
async def export_workspace_bundle(project_name: str = Query(default="")) -> StreamingResponse:
    if project_name.strip():
        config = read_config()
        config["project_name"] = normalize_project_name(project_name)
        config["project_updated_at"] = utc_now()
        write_config(config)
    project = current_project_name()
    payload = build_workspace_zip(project)
    filename = f"{project_filename_slug(project)}-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.zip"
    return StreamingResponse(
        iter([payload]),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/import/workspace")
async def import_workspace_bundle(file: UploadFile = File(...)) -> dict[str, Any]:
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded workspace bundle is empty.")
    try:
        imported = restore_workspace_zip(content)
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400, detail="Uploaded file is not a valid workspace zip bundle.") from exc
    project_name = current_project_name()
    return {"ok": True, "imported": imported, "project": {"name": project_name, "slug": project_slug(project_name)}}


@app.post("/api/apply-ranges")
async def apply_ranges(payload: ApplyRangesPayload) -> dict[str, Any]:
    updates = payload.updates
    if not updates:
        raise HTTPException(status_code=400, detail="No range updates were provided.")
    app_config = read_config()
    requested_live_edit = payload.requested_mode == "live_edit"
    if not requested_live_edit:
        instructions = build_manual_instructions(updates, scope=payload.scope, mode="read_only")
        INSTRUCTIONS_PATH.write_text(json.dumps(instructions, indent=2), encoding="utf-8")
        document = save_instruction_document(instructions, payload.visualizations)
        return {
            "ok": True,
            "mode": "read_only",
            "instruction_count": len(instructions["steps"]),
            "instructions": instructions,
            "document": document,
        }
    if not app_config.get("live_edit_enabled", False):
        raise HTTPException(status_code=409, detail="Live editing is disabled. Switch to LIVE EDITING before sending Admin API range changes.")
    config = saved_connection("admin")
    instructions = build_manual_instructions(updates, scope=payload.scope, mode="live_applied")
    try:
        result = AdminApiClient(config).apply_range_updates(updates)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    rows = normalize_admin_segments(result.get("segments", []))
    ADMIN_SEGMENTS_PATH.write_text(
        json.dumps({"collected_at": utc_now(), "count": len(rows), "raw": result.get("segments", []), "segments": rows}, indent=2),
        encoding="utf-8",
    )
    INSTRUCTIONS_PATH.write_text(json.dumps(instructions, indent=2), encoding="utf-8")
    document = save_instruction_document(instructions, payload.visualizations)
    return {
        "ok": True,
        "applied": result["applied"],
        "segments": len(rows),
        "zero_ranges": [row for row in rows if not row.get("ranges")],
        "document": document,
    }


@app.get("/api/download/scrm-offline-host-ip-collector.py")
@app.get("/api/download/offline-host-collector.py")
async def download_offline_host_collector() -> StreamingResponse:
    return StreamingResponse(
        iter([offline_host_collector_script()]),
        media_type="text/x-python",
        headers={"Content-Disposition": 'attachment; filename="scrm_offline_host_ip_collector.py"'},
    )


@app.get("/api/export/manual-instructions.csv")
async def export_manual_instructions() -> StreamingResponse:
    instructions = read_json(INSTRUCTIONS_PATH, {"steps": []})

    return StreamingResponse(
        iter([instructions_to_csv(instructions)]),
        media_type="text/csv",
        headers={"Content-Disposition": 'attachment; filename="manual-range-change-instructions.csv"'},
    )


@app.get("/api/documents")
async def list_generated_documents() -> dict[str, Any]:
    ensure_latest_instruction_document()
    return {"documents": list_documents()}


@app.get("/api/documents/{document_id}")
async def download_generated_document(document_id: str) -> FileResponse:
    path = document_path(document_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Generated document was not found.")
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if path.suffix.lower() == ".docx" else "text/csv"
    return FileResponse(path, media_type=media_type, filename=path.name)


@app.delete("/api/documents/{document_id}")
async def delete_generated_document(document_id: str) -> dict[str, Any]:
    path = document_path(document_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Generated document was not found.")
    meta_path = path.with_suffix(".json")
    csv_path = path.with_suffix(".csv")
    path.unlink()
    if meta_path.exists():
        meta_path.unlink()
    if csv_path.exists():
        csv_path.unlink()
    return {"ok": True, "deleted": document_id, "documents": list_documents()}


@app.get("/api/export/zero-ranges.csv")
async def export_zero_ranges() -> StreamingResponse:
    analysis = build_analysis()
    rows = analysis["stages"]["zero_ranges"]

    def iter_csv():
        from io import StringIO

        stream = StringIO()
        writer = csv.writer(stream)
        writer.writerow(["name", "path", "used", "policy_reference_count", "child_count"])
        yield stream.getvalue()
        stream.seek(0)
        stream.truncate(0)
        for row in rows:
            writer.writerow([row["name"], row["path"], row["used"], row["policy_reference_count"], row["child_count"]])
            yield stream.getvalue()
            stream.seek(0)
            stream.truncate(0)

    return StreamingResponse(
        iter_csv(),
        media_type="text/csv",
        headers={"Content-Disposition": 'attachment; filename="zero-range-segments.csv"'},
    )


app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")


def build_analysis() -> dict[str, Any]:
    policies = load_parsed_artifact("policies")
    xml_segments = load_parsed_artifact("segments")
    host_snapshot = read_json(HOSTS_PATH, {"hosts": []})
    admin_snapshot = read_json(ADMIN_SEGMENTS_PATH, {"segments": []})
    live_segments = admin_snapshot.get("segments") or []
    if not live_segments:
        live_segments = xml_segments.get("segments") or []

    usage = build_policy_segment_usage(policies.get("policies") or [], xml_segments.get("segments") or [], live_segments)
    host_ips = [host["ip"] for host in host_snapshot.get("hosts") or [] if host.get("ip")]
    conflicts = detect_conflicts(live_segments, usage, host_ips)
    mapping = build_segment_policy_mapping(live_segments, usage, conflicts, policies.get("policies") or [])
    zero_ranges = [
        {
            **usage.get(segment["key"], {}),
            "key": segment["key"],
            "name": segment["name"],
            "path": segment["path"],
            "ranges": segment.get("ranges") or [],
            "child_count": segment.get("child_count", 0),
        }
        for segment in live_segments
        if not segment.get("ranges")
    ]
    return {
        "ready": bool(policies.get("policies") and xml_segments.get("segments") and live_segments),
        "summary": {
            "policies": len(policies.get("policies") or []),
            "xml_segments": len(xml_segments.get("segments") or []),
            "segments": len(live_segments),
            "live_segments": len(live_segments),
            "host_ips": len(host_ips),
            "conflicts": sum(len(rows) for rows in conflicts.values()),
            "zero_ranges": len(zero_ranges),
        },
        "stages": {
            "used_wins": conflicts["used_wins"],
            "ownership": conflicts["ownership"],
            "live_decision": conflicts["live_decision"],
            "lower_review": conflicts["lower_review"],
            "zero_ranges": zero_ranges,
        },
        "mapping": mapping,
    }


def detect_conflicts(segments: list[dict[str, Any]], usage: dict[str, dict[str, Any]], host_ips: list[str]) -> dict[str, list[dict[str, Any]]]:
    host_ints = []
    for ip in host_ips:
        try:
            host_ints.append(int(ipaddress.IPv4Address(ip)))
        except Exception:
            continue
    ranges = []
    for segment in segments:
        segment_usage = usage.get(segment["key"], {})
        for range_value in segment.get("ranges") or []:
            parsed = parse_ip_range(range_value)
            if parsed:
                ranges.append({"segment": {**segment, **segment_usage}, "range": range_value, **parsed})
    ranges.sort(key=lambda row: (row["start"], row["end"], row["segment"]["path"]))
    output = {"used_wins": [], "ownership": [], "live_decision": [], "lower_review": []}
    seen = set()
    for left_index, left in enumerate(ranges):
        for right in ranges[left_index + 1 :]:
            if right["start"] > left["end"]:
                break
            if left["segment"]["key"] == right["segment"]["key"]:
                continue
            start = max(left["start"], right["start"])
            end = min(left["end"], right["end"])
            if start > end:
                continue
            key = "::".join(sorted([left["segment"]["key"], right["segment"]["key"]])) + f"::{start}-{end}"
            if key in seen:
                continue
            seen.add(key)
            live_ips = [str(ipaddress.IPv4Address(ip)) for ip in host_ints if start <= ip <= end]
            left_used = bool(left["segment"].get("used"))
            right_used = bool(right["segment"].get("used"))
            row = {
                "id": key,
                "overlap_range": format_ip_range(start, end),
                "ip_count": end - start + 1,
                "live_host_count": len(live_ips),
                "live_ips": live_ips,
                "left": conflict_segment(left),
                "right": conflict_segment(right),
                "left_range": left["range"],
                "right_range": right["range"],
            }
            if left_used != right_used:
                target = right if left_used else left
                row["default_update"] = range_update_for(target, start, end, "Policy Usage Wins", "Policy-used segment wins over not-used segment")
                output["used_wins"].append(row)
            elif not left_used and not right_used:
                keep_left = authoritative_segment(left["segment"], right["segment"])["key"] == left["segment"]["key"]
                row["default_keep"] = "left" if keep_left else "right"
                output["ownership"].append(row)
            elif live_ips:
                output["live_decision"].append(row)
            else:
                output["lower_review"].append(row)
    for rows in output.values():
        rows.sort(key=lambda row: (-row["live_host_count"], -row["ip_count"], row["overlap_range"]))
    return output


def build_manual_instructions(updates: list[RangeUpdate], scope: str = "", mode: str = "read_only") -> dict[str, Any]:
    project = current_project_name()
    segments = current_segments_for_analysis()
    by_key = {row.get("key", ""): row for row in segments}
    by_path = {_normalize_path(row.get("path", "")): row for row in segments if row.get("path")}
    xml_descriptions = segment_description_lookup()
    grouped: dict[str, list[RangeUpdate]] = {}
    for update in updates:
        key = update.segment_key or _normalize_path(update.path) or f"{update.name}:{update.source_range}"
        grouped.setdefault(key, []).append(update)

    steps = []
    step_number = 1
    for group_updates in grouped.values():
        first = group_updates[0]
        segment = by_key.get(first.segment_key) or by_path.get(_normalize_path(first.path)) or {}
        current_ranges = list(segment.get("ranges") or [])
        resulting_ranges = subtract_updates_from_ranges(current_ranges, group_updates) if current_ranges else []
        for update in group_updates:
            segment_name = segment.get("name") or update.name
            segment_path = segment.get("path") or update.path
            description = (
                str(segment.get("description") or "").strip()
                or xml_descriptions.get(segment.get("key", ""))
                or xml_descriptions.get(f"{segment_name}::{_normalize_path(segment_path)}", "")
            )
            instruction = (
                f"In Segment Manager, edit '{segment_name}' under '{segment_path}'. "
                f"Remove overlap {update.remove_range} from source range {update.source_range}. "
                "Do not rename, move, create, or delete the segment for this step."
            )
            steps.append(
                {
                    "step": step_number,
                    "stage": update.stage,
                    "name": segment_name,
                    "path": segment_path,
                    "description": description,
                    "source_range": update.source_range,
                    "remove_range": update.remove_range,
                    "current_ranges": current_ranges,
                    "resulting_ranges": resulting_ranges,
                    "reason": update.reason,
                    "instruction": instruction,
                }
            )
            step_number += 1
    summary_label = "live range change record" if mode == "live_applied" else "manual range change instruction"
    return {
        "generated_at": utc_now(),
        "project_name": project,
        "mode": mode,
        "scope": scope,
        "summary": f"{len(steps)} {summary_label}{'s' if len(steps) != 1 else ''} generated.",
        "steps": steps,
    }


def instructions_to_csv(instructions: dict[str, Any]) -> str:
    from io import StringIO

    stream = StringIO()
    writer = csv.writer(stream)
    writer.writerow(["project", "step", "stage", "segment", "path", "source_range", "remove_range", "resulting_ranges", "reason", "instruction"])
    project = instructions.get("project_name") or current_project_name()
    for row in instructions.get("steps") or []:
        writer.writerow(
            [
                project,
                row.get("step", ""),
                row.get("stage", ""),
                row.get("name", ""),
                row.get("path", ""),
                row.get("source_range", ""),
                row.get("remove_range", ""),
                "; ".join(row.get("resulting_ranges") or []),
                row.get("reason", ""),
                row.get("instruction", ""),
            ]
        )
    return stream.getvalue()


def workspace_file_map() -> dict[str, Path]:
    return {
        "uploads/policies.xml": POLICIES_PATH,
        "uploads/segments.xml": SEGMENTS_PATH,
        "uploads/policies.json": UPLOAD_DIR / "policies.json",
        "uploads/segments.json": UPLOAD_DIR / "segments.json",
        "snapshots/hosts.json": HOSTS_PATH,
        "snapshots/admin_segments.json": ADMIN_SEGMENTS_PATH,
        "snapshots/manual_instructions.json": INSTRUCTIONS_PATH,
    }


def clear_workspace_files(include_documents: bool) -> list[str]:
    targets = list(workspace_file_map().values())
    if include_documents:
        targets.extend(path for path in DOCUMENTS_DIR.glob("*") if path.is_file())
    removed = []
    for path in targets:
        if path.exists():
            path.unlink()
            removed.append(str(path.relative_to(DATA_DIR)))
    return removed


def build_workspace_zip(project_name: str | None = None) -> bytes:
    ensure_latest_instruction_document()
    project = normalize_project_name(project_name or current_project_name())
    buffer = io.BytesIO()
    manifest = {
        "app": "Segment Conflict Resolution Management",
        "format": "segment-conflict-workspace",
        "version": 2,
        "project_name": project,
        "project_slug": project_slug(project),
        "created_at": utc_now(),
        "note": "This bundle contains loaded artifacts, snapshots, generated documents, and sanitized API metadata. Passwords are intentionally excluded.",
        "files": [],
    }
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        for arcname, path in workspace_file_map().items():
            if path.exists():
                archive.write(path, arcname)
                manifest["files"].append(arcname)
        sanitized_config = sanitized_workspace_config()
        if sanitized_config:
            archive.writestr("config/api-metadata.json", json.dumps(sanitized_config, indent=2))
            manifest["files"].append("config/api-metadata.json")
        for path in sorted(DOCUMENTS_DIR.glob("*")):
            if path.is_file() and path.suffix.lower() in {".csv", ".docx", ".json"}:
                arcname = f"Documents/{path.name}"
                archive.write(path, arcname)
                manifest["files"].append(arcname)
        archive.writestr("manifest.json", json.dumps(manifest, indent=2))
    return buffer.getvalue()


def restore_workspace_zip(content: bytes) -> list[str]:
    imported: list[str] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        members = [(member, workspace_target_for_zip_member(member.filename)) for member in archive.infolist()]
        for member, target in members:
            if target and member.file_size > 50 * 1024 * 1024:
                raise HTTPException(status_code=400, detail=f"Workspace member is too large: {member.filename}")
        clear_workspace_files(include_documents=True)
        restored_config = sanitized_config_from_archive(archive)
        if restored_config:
            write_config(restored_config)
            imported.append(str(CONFIG_PATH.relative_to(DATA_DIR)))
        for member in archive.infolist():
            target = workspace_target_for_zip_member(member.filename)
            if not target:
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(archive.read(member))
            imported.append(str(target.relative_to(DATA_DIR)))
    return imported


def sanitized_workspace_config() -> dict[str, Any]:
    config = read_config()
    sanitized: dict[str, Any] = {
        "project_name": current_project_name(config),
        "project_updated_at": config.get("project_updated_at", ""),
    }
    for kind in ("web", "admin"):
        row = config.get(kind) or {}
        if row.get("base_url") or row.get("username"):
            sanitized[kind] = {
                "base_url": row.get("base_url", ""),
                "username": row.get("username", ""),
                "password": "",
                "verify_tls": bool(row.get("verify_tls", False)),
                "updated_at": row.get("updated_at", ""),
            }
    if "live_edit_enabled" in config:
        sanitized["live_edit_enabled"] = bool(config.get("live_edit_enabled", False))
        sanitized["protection_updated_at"] = config.get("protection_updated_at", "")
    return sanitized


def sanitized_config_from_archive(archive: zipfile.ZipFile) -> dict[str, Any]:
    try:
        raw = archive.read("config/api-metadata.json")
    except KeyError:
        return {}
    try:
        source = json.loads(raw.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError):
        raise HTTPException(status_code=400, detail="Workspace API metadata is not valid JSON.")
    restored: dict[str, Any] = {}
    restored["project_name"] = normalize_project_name(str(source.get("project_name") or DEFAULT_PROJECT_NAME))
    restored["project_updated_at"] = str(source.get("project_updated_at") or utc_now())
    for kind in ("web", "admin"):
        row = source.get(kind) or {}
        if row.get("base_url") or row.get("username"):
            restored[kind] = {
                "base_url": str(row.get("base_url", "")),
                "username": str(row.get("username", "")),
                "password": "",
                "verify_tls": bool(row.get("verify_tls", False)),
                "updated_at": str(row.get("updated_at", "")),
            }
    if "live_edit_enabled" in source:
        restored["live_edit_enabled"] = bool(source.get("live_edit_enabled", False))
        restored["protection_updated_at"] = str(source.get("protection_updated_at", ""))
    return restored


def workspace_target_for_zip_member(name: str) -> Path | None:
    normalized = name.replace("\\", "/").strip("/")
    if not normalized or normalized.endswith("/") or normalized == "manifest.json":
        return None
    fixed = workspace_file_map().get(normalized)
    if fixed:
        return fixed
    if normalized.lower().startswith("documents/"):
        filename = normalized.split("/", 1)[1]
        if "/" not in filename and re.fullmatch(r"[A-Za-z0-9._-]+\\.(csv|docx|json)", filename, flags=re.IGNORECASE):
            return DOCUMENTS_DIR / filename
    return None


def save_instruction_document(instructions: dict[str, Any], visualizations: list[VisualizationPayload] | None = None) -> dict[str, Any]:
    generated_at = instructions.get("generated_at") or utc_now()
    timestamp = re.sub(r"[^0-9]", "", generated_at)[:14] or datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    project = normalize_project_name(str(instructions.get("project_name") or current_project_name()))
    steps = instructions.get("steps") or []
    stage_names = sorted({str(step.get("stage") or "manual") for step in steps})
    stage_slug = safe_slug("-".join(stage_names) or "manual-range-changes")
    base_name = f"{project_filename_slug(project)}-{timestamp}-{stage_slug}-recommendations"
    path = unique_document_path(base_name, ".docx")
    csv_path = path.with_suffix(".csv")
    csv_path.write_text(instructions_to_csv(instructions), encoding="utf-8")
    write_instruction_docx(path, instructions, visualizations or [])
    metadata = {
        "id": path.name,
        "filename": path.name,
        "project_name": project,
        "created_at": generated_at,
        "summary": instructions.get("summary", ""),
        "step_count": len(steps),
        "stages": stage_names,
        "size": path.stat().st_size,
        "scope": instructions.get("scope", ""),
        "mode": instructions.get("mode", "read_only"),
        "csv_filename": csv_path.name,
    }
    path.with_suffix(".json").write_text(json.dumps(metadata, indent=2), encoding="utf-8")
    return metadata


def write_instruction_docx(path: Path, instructions: dict[str, Any], visualizations: list[VisualizationPayload]) -> None:
    project = normalize_project_name(str(instructions.get("project_name") or current_project_name()))
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    configure_docx_styles(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.add_run(f"{project} - Segment Conflict Resolution Recommendations")

    subtitle = doc.add_paragraph()
    subtitle.style = "Subtitle"
    mode_label = "Live changes applied" if instructions.get("mode") == "live_applied" else "Read-only admin instructions"
    subtitle.add_run(f"{mode_label} | Generated {instructions.get('generated_at', utc_now())}")
    add_note(doc, "Project", project)

    scope = str(instructions.get("scope") or "").strip()
    if scope:
        add_note(doc, "Scope", scope)

    steps = instructions.get("steps") or []
    stages = sorted({str(step.get("stage") or "manual") for step in steps})
    add_note(
        doc,
        "Decision Summary",
        f"{len(steps)} range change decision{'s' if len(steps) != 1 else ''} across {len(stages)} workflow stage{'s' if len(stages) != 1 else ''}. "
        "The API boundary is range updates only: no segment create, rename, move, or delete action is included in this document.",
    )

    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    headers = ["Steps", "Affected Segments", "Affected Ranges", "Mode"]
    for index, header in enumerate(headers):
        cell = summary.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, "EAF2FF")
    row = summary.add_row().cells
    row[0].text = str(len(steps))
    row[1].text = str(len({step.get("path") or step.get("name") for step in steps}))
    row[2].text = str(len({step.get("remove_range") for step in steps}))
    row[3].text = mode_label

    visualization_by_range = visualizations_by_range(visualizations)
    doc.add_heading("Resolution Decisions", level=1)
    if not steps:
        doc.add_paragraph("No range change decisions were generated.")
    for step in steps:
        add_resolution_step(doc, step, visualization_by_range.get(str(step.get("remove_range") or "")))

    doc.add_heading("Operator Notes", level=1)
    notes = [
        "Validate each policy owner decision before applying manual changes to production.",
        "Retained segment IDs are preserved because the recommendation removes only ranges from existing segment objects.",
        "If a segment becomes zero-range after cleanup, review it separately before deleting it from the Segment Manager or from a cleaned XML export.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    doc.save(path)


def configure_docx_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (("Title", 20, "111827"), ("Subtitle", 10.5, "64748B"), ("Heading 1", 15, "1E3A8A"), ("Heading 2", 12, "111827")):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(8 if name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(5)


def add_note(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    paragraph.add_run(text)


def add_resolution_step(doc: Document, step: dict[str, Any], visualization: dict[str, Any] | None) -> None:
    step_number = step.get("step", "")
    segment_name = step.get("name") or "Unnamed segment"
    remove_range = step.get("remove_range") or "Unknown range"
    doc.add_heading(f"Decision {step_number}: Remove {remove_range} from {segment_name}", level=2)

    if step.get("description"):
        add_note(doc, "Segment XML description", str(step.get("description")))
    else:
        add_note(doc, "Segment XML description", "No segment description was present in the uploaded Segments.xml export.")

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Workflow stage", stage_display_name(str(step.get("stage") or ""))),
        ("Segment path", step.get("path") or ""),
        ("Source range", step.get("source_range") or ""),
        ("Range to remove", remove_range),
        ("Current ranges", "; ".join(step.get("current_ranges") or []) or "None recorded"),
        ("Resulting ranges", "; ".join(step.get("resulting_ranges") or []) or "None"),
        ("Decision reason", step.get("reason") or ""),
        ("Operator instruction", step.get("instruction") or ""),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        shade_cell(cells[0], "F8FAFC")
    if visualization:
        doc.add_paragraph("Resolution visualization", style=None).runs[0].bold = True
        image = visualization_image_bytes(visualization)
        if image:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image, width=Inches(6.8))
            caption = doc.add_paragraph(str(visualization.get("title") or remove_range))
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph("Visualization was unavailable for this resolution.")


def visualizations_by_range(visualizations: list[VisualizationPayload]) -> dict[str, dict[str, Any]]:
    output: dict[str, dict[str, Any]] = {}
    for item in visualizations:
        row = item.model_dump()
        key = str(row.get("range") or "").strip()
        if key and row.get("png_data_url"):
            output.setdefault(key, row)
    return output


def visualization_image_bytes(visualization: dict[str, Any]) -> io.BytesIO | None:
    data_url = str(visualization.get("png_data_url") or "")
    if "," not in data_url or not data_url.startswith("data:image/png"):
        return None
    try:
        payload = base64.b64decode(data_url.split(",", 1)[1], validate=True)
    except (ValueError, binascii.Error):
        return None
    return io.BytesIO(payload)


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def stage_display_name(stage: str) -> str:
    return {
        "used_wins": "Policy Usage Wins",
        "ownership": "Ownership Decisions",
        "live_decision": "Admin or Policy Owner Decision",
        "lower_review": "Lower Priority Review",
    }.get(stage, stage or "Manual")


def ensure_latest_instruction_document() -> None:
    instructions = read_json(INSTRUCTIONS_PATH, {})
    steps = instructions.get("steps") or []
    generated_at = instructions.get("generated_at")
    if not steps or not generated_at:
        return
    for meta_path in DOCUMENTS_DIR.glob("*.json"):
        metadata = read_json(meta_path, {})
        candidate = DOCUMENTS_DIR / str(metadata.get("filename") or "")
        if metadata.get("created_at") == generated_at and int(metadata.get("step_count") or 0) == len(steps) and candidate.suffix.lower() == ".docx" and candidate.exists():
            return
    save_instruction_document(instructions)


def list_documents() -> list[dict[str, Any]]:
    documents = []
    for path in DOCUMENTS_DIR.glob("*.docx"):
        metadata = read_json(path.with_suffix(".json"), {})
        stat = path.stat()
        documents.append(
            {
                "id": path.name,
                "filename": path.name,
                "project_name": metadata.get("project_name") or current_project_name(),
                "created_at": metadata.get("created_at") or datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
                "summary": metadata.get("summary", ""),
                "step_count": metadata.get("step_count", 0),
                "stages": metadata.get("stages", []),
                "scope": metadata.get("scope", ""),
                "mode": metadata.get("mode", ""),
                "size": stat.st_size,
            }
        )
    documents.sort(key=lambda item: item.get("created_at") or "", reverse=True)
    return documents


def document_path(document_id: str) -> Path:
    if "/" in document_id or "\\" in document_id or not document_id.lower().endswith((".docx", ".csv")):
        raise HTTPException(status_code=400, detail="Invalid document id.")
    path = (DOCUMENTS_DIR / document_id).resolve()
    if path.parent != DOCUMENTS_DIR.resolve():
        raise HTTPException(status_code=400, detail="Invalid document path.")
    return path


def unique_document_path(base_name: str, suffix: str = ".docx") -> Path:
    path = DOCUMENTS_DIR / f"{base_name}{suffix}"
    if not path.exists():
        return path
    version = 2
    while True:
        candidate = DOCUMENTS_DIR / f"{base_name}-v{version}{suffix}"
        if not candidate.exists():
            return candidate
        version += 1


def safe_slug(value: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-").lower()
    return slug[:80] or "recommendations"


def normalize_project_name(value: str) -> str:
    name = re.sub(r"\s+", " ", str(value or "").strip())
    return name[:120] or DEFAULT_PROJECT_NAME


def current_project_name(config: dict[str, Any] | None = None) -> str:
    source = config if config is not None else read_config()
    return normalize_project_name(str(source.get("project_name") or DEFAULT_PROJECT_NAME))


def project_filename_slug(value: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", normalize_project_name(value)).strip("._-")
    slug = re.sub(r"_+", "_", slug)
    return slug[:80] or "Segment_Conflict_Workspace"


def project_slug(value: str) -> str:
    return safe_slug(normalize_project_name(value))


def current_segments_for_analysis() -> list[dict[str, Any]]:
    admin_snapshot = read_json(ADMIN_SEGMENTS_PATH, {"segments": []})
    live_segments = admin_snapshot.get("segments") or []
    if live_segments:
        return live_segments
    xml_segments = load_parsed_artifact("segments")
    return xml_segments.get("segments") or []


def segment_description_lookup() -> dict[str, str]:
    lookup: dict[str, str] = {}
    xml_segments = load_parsed_artifact("segments").get("segments") or []
    if SEGMENTS_PATH.exists():
        try:
            xml_segments = parse_segments_xml(SEGMENTS_PATH.read_bytes()).get("segments") or xml_segments
        except Exception:
            pass
    admin_segments = read_json(ADMIN_SEGMENTS_PATH, {"segments": []}).get("segments") or []
    for rowset in (xml_segments, admin_segments):
        for source in rowset:
            if not isinstance(source, dict):
                continue
            key = str(source.get("key") or "")
            description = str(source.get("description") or "").strip()
            if key and description:
                lookup[key] = description
            name = str(source.get("name") or "")
            path = str(source.get("path") or "")
            if name and path and description:
                lookup[f"{name}::{_normalize_path(path)}"] = description
    return lookup


def conflict_segment(row: dict[str, Any]) -> dict[str, Any]:
    segment = row["segment"]
    return {
        "key": segment.get("key", ""),
        "name": segment.get("name", ""),
        "path": segment.get("path", ""),
        "depth": segment.get("depth", 0),
        "child_count": segment.get("child_count", 0),
        "ranges": segment.get("ranges") or [],
        "used": bool(segment.get("used")),
        "policy_reference_count": segment.get("policy_reference_count", 0),
        "policy_references": segment.get("policy_references") or [],
        "direct_used": bool(segment.get("direct_used")),
        "used_reason": segment.get("used_reason", ""),
        "range": row["range"],
    }


def range_update_for(row: dict[str, Any], start: int, end: int, stage: str, reason: str) -> dict[str, str]:
    segment = row["segment"]
    return {
        "segment_key": segment.get("key", ""),
        "name": segment.get("name", ""),
        "path": segment.get("path", ""),
        "source_range": row["range"],
        "remove_range": format_ip_range(start, end),
        "reason": reason,
        "stage": stage,
    }


def build_policy_segment_usage(policies: list[dict[str, Any]], xml_segments: list[dict[str, Any]], live_segments: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    xml_by_id = {str(row.get("id")): row for row in xml_segments if row.get("id")}
    name_counts: dict[str, int] = {}
    for row in xml_segments:
        name_counts[row.get("name", "").casefold()] = name_counts.get(row.get("name", "").casefold(), 0) + 1
    xml_by_unique_name = {row.get("name", "").casefold(): row for row in xml_segments if name_counts.get(row.get("name", "").casefold()) == 1}
    direct_paths: set[str] = set()
    refs_by_path: dict[str, list[dict[str, str]]] = {}
    for policy in policies:
        for ref in policy.get("segment_refs", []):
            segment = None
            if ref.get("id") and ref["id"] in xml_by_id:
                segment = xml_by_id[ref["id"]]
            elif ref.get("name") and ref["name"].casefold() in xml_by_unique_name:
                segment = xml_by_unique_name[ref["name"].casefold()]
            if not segment:
                continue
            path_key = _normalize_path(segment.get("path", ""))
            direct_paths.add(path_key)
            refs_by_path.setdefault(path_key, []).append(
                {
                    "policy": policy.get("name", ""),
                    "source": ref.get("source", ""),
                    "segment_path": segment.get("path", ""),
                    "inherited": False,
                }
            )

    usage: dict[str, dict[str, Any]] = {}
    for segment in live_segments:
        path_key = _normalize_path(segment.get("path", ""))
        direct = path_key in direct_paths
        inherited_from = max((path for path in direct_paths if path_key.startswith(f"{path}/")), key=len, default="")
        used = direct or bool(inherited_from)
        references = [*refs_by_path.get(path_key, [])]
        if inherited_from:
            for reference in refs_by_path.get(inherited_from, []):
                references.append({**reference, "source": f"{reference.get('source', '')} via parent segment".strip(), "inherited": True})
        usage[segment["key"]] = {
            "used": used,
            "direct_used": direct,
            "used_reason": "direct policy segment" if direct else "under policy-used parent" if inherited_from else "",
            "policy_reference_count": len(references),
            "policy_references": references,
        }
    return usage


def build_segment_policy_mapping(
    live_segments: list[dict[str, Any]],
    usage: dict[str, dict[str, Any]],
    conflicts: dict[str, list[dict[str, Any]]],
    policies: list[dict[str, Any]],
) -> dict[str, Any]:
    conflicting_ranges_by_segment: dict[str, set[str]] = {}
    conflict_categories_by_segment: dict[str, set[str]] = {}
    for stage_key, rows in conflicts.items():
        for row in rows:
            overlap = str(row.get("overlap_range") or "")
            for side in ("left", "right"):
                segment = row.get(side) or {}
                segment_key = str(segment.get("key") or "")
                if not segment_key:
                    continue
                if overlap:
                    conflicting_ranges_by_segment.setdefault(segment_key, set()).add(overlap)
                conflict_categories_by_segment.setdefault(segment_key, set()).add(stage_key)

    policy_lookup = {str(policy.get("name") or ""): policy for policy in policies}
    policy_segments: dict[str, dict[str, Any]] = {}
    segment_rows: list[dict[str, Any]] = []
    for segment in live_segments:
        segment_key = str(segment.get("key") or "")
        segment_usage = usage.get(segment_key, {})
        references = segment_usage.get("policy_references") or []
        conflict_ranges = sorted(conflicting_ranges_by_segment.get(segment_key, set()), key=lambda value: [_ip_sort_token(value), value])
        categories = sorted(conflict_categories_by_segment.get(segment_key, set()))
        row = {
            "key": segment_key,
            "name": segment.get("name", ""),
            "path": segment.get("path", ""),
            "depth": segment.get("depth", 0),
            "ranges": segment.get("ranges") or [],
            "child_count": segment.get("child_count", 0),
            "used": bool(segment_usage.get("used")),
            "direct_used": bool(segment_usage.get("direct_used")),
            "used_reason": segment_usage.get("used_reason", ""),
            "policy_reference_count": int(segment_usage.get("policy_reference_count") or 0),
            "policy_references": references,
            "has_conflicts": bool(conflict_ranges),
            "conflict_range_count": len(conflict_ranges),
            "conflicting_ranges": conflict_ranges,
            "categories": categories,
        }
        segment_rows.append(row)
        for reference in references:
            policy_name = str(reference.get("policy") or "Unnamed policy").strip() or "Unnamed policy"
            if policy_name not in policy_segments:
                policy = policy_lookup.get(policy_name, {})
                policy_segments[policy_name] = {
                    "policy": policy_name,
                    "folder": policy.get("folder", ""),
                    "enabled": bool(policy.get("enabled", True)),
                    "sources": set(),
                    "segments": {},
                }
            stored = policy_segments[policy_name]
            if reference.get("source"):
                stored["sources"].add(str(reference.get("source")))
            stored["segments"][segment_key] = {
                "key": segment_key,
                "name": segment.get("name", ""),
                "path": segment.get("path", ""),
                "used": bool(segment_usage.get("used")),
                "has_conflicts": bool(conflict_ranges),
                "conflict_range_count": len(conflict_ranges),
                "conflicting_ranges": conflict_ranges,
            }

    policy_rows = []
    for policy in policy_segments.values():
        segments = sorted(
            policy["segments"].values(),
            key=lambda segment: (not segment.get("has_conflicts"), str(segment.get("path", "")).casefold(), str(segment.get("name", "")).casefold()),
        )
        conflicting_segment_count = sum(1 for segment in segments if segment.get("has_conflicts"))
        policy_rows.append(
            {
                "policy": policy["policy"],
                "folder": policy.get("folder", ""),
                "enabled": policy.get("enabled", True),
                "sources": sorted(policy["sources"], key=lambda value: value.casefold()),
                "segments": segments,
                "segment_count": len(segments),
                "conflicting_segment_count": conflicting_segment_count,
                "mapping_state": "conflicting_segments" if conflicting_segment_count else "clean_segments",
            }
        )

    for policy in policies:
        policy_name = str(policy.get("name") or "Unnamed policy").strip() or "Unnamed policy"
        if policy_name in policy_segments:
            continue
        policy_rows.append(
            {
                "policy": policy_name,
                "folder": policy.get("folder", ""),
                "enabled": bool(policy.get("enabled", True)),
                "sources": [],
                "segments": [],
                "segment_count": 0,
                "conflicting_segment_count": 0,
                "mapping_state": "no_segments",
            }
        )

    segment_rows.sort(key=lambda segment: (str(segment.get("path", "")).casefold(), str(segment.get("name", "")).casefold()))
    policy_rows.sort(key=lambda policy: (-int(policy.get("conflicting_segment_count") or 0), str(policy.get("policy", "")).casefold()))
    return {
        "segments": segment_rows,
        "policies": policy_rows,
        "summary": {
            "segments": len(segment_rows),
            "policies": len(policy_rows),
            "conflicting_segments": sum(1 for segment in segment_rows if segment.get("has_conflicts")),
            "clean_segments": sum(1 for segment in segment_rows if not segment.get("has_conflicts")),
            "policies_without_segments": sum(1 for policy in policy_rows if policy.get("mapping_state") == "no_segments"),
            "policies_linked_to_conflicting_segments": sum(1 for policy in policy_rows if policy.get("mapping_state") == "conflicting_segments"),
        },
    }


def parse_policies_xml(content: bytes) -> dict[str, Any]:
    root = parse_xml(content)
    if local_name(root.tag) != "POLICY_FOLDER":
        raise HTTPException(status_code=400, detail="Policies XML must start with POLICY_FOLDER.")
    policies = []

    def walk(folder: ET.Element, path: list[str]) -> None:
        folder_name = folder.attrib.get("NAME", "Policy Folders")
        folder_path = [*path, folder_name]
        policies_node = direct_child(folder, "POLICIES")
        for policy in direct_children(policies_node, "POLICY"):
            rule = direct_child(policy, "RULE")
            if rule is None:
                continue
            refs = []
            refs.extend(rule_segment_refs(rule, "Policy segment"))
            for inner in iter_tag(rule, "INNER_RULE"):
                refs.extend(rule_segment_refs(inner, "Sub-rule segment"))
            policies.append(
                {
                    "id": rule.attrib.get("ID", ""),
                    "name": rule.attrib.get("NAME", "Unnamed policy"),
                    "enabled": bool_attr(rule.attrib.get("ENABLED"), True),
                    "folder": " / ".join(folder_path),
                    "segment_refs": unique_refs(refs),
                }
            )
        for child in direct_children(folder, "POLICY_FOLDER"):
            walk(child, folder_path)

    walk(root, [])
    return {"summary": {"policies": len(policies), "segment_refs": sum(len(policy["segment_refs"]) for policy in policies)}, "policies": policies}


def parse_segments_xml(content: bytes) -> dict[str, Any]:
    root = parse_xml(content)
    if local_name(root.tag) != "GROUP":
        raise HTTPException(status_code=400, detail="Segments XML must start with GROUP.")
    rows = []

    def walk(node: ET.Element, path: list[str], root_node: bool = False) -> None:
        name = node.attrib.get("NAME", "Segments")
        display = "Segments" if root_node else name
        current_path = [*path, display]
        children = direct_children(node, "GROUP")
        ranges = [child.attrib.get("RANGE", "") for child in direct_children(node, "RANGES") if child.attrib.get("RANGE")]
        if not root_node:
            segment_path = " / ".join(current_path)
            rows.append(
                {
                    "id": node.attrib.get("SEGMENT_ID", node.attrib.get("ID", "")),
                    "key": f"{name}::{_normalize_path(segment_path)}",
                    "name": name,
                    "path": segment_path,
                    "description": node.attrib.get("DESCRIPTION", ""),
                    "depth": len(current_path) - 1,
                    "ranges": ranges,
                    "child_count": len(children),
                }
            )
        for child in children:
            walk(child, current_path)

    walk(root, [], root_node=True)
    return {"summary": {"segments": len(rows), "ranges": sum(len(row["ranges"]) for row in rows)}, "segments": rows}


def normalize_admin_segments(payload: dict[str, Any] | list[Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for root in admin_roots(payload):
        if isinstance(root, dict):
            walk_admin_node(root, [], rows, root_node=True)
    return rows


def walk_admin_node(node: dict[str, Any], path: list[str], rows: list[dict[str, Any]], root_node: bool = False) -> None:
    name = str(pick(node, "name", "Name") or "Unnamed segment").strip()
    display = "Segments" if root_node and name == "*root" else name
    current_path = [*path, display] if display else path
    children = first_list(node, "nodes", "Nodes", "children", "Children")
    ranges = [str(item).strip() for item in first_list(node, "ranges", "Ranges") if str(item).strip()]
    if not root_node:
        segment_path = " / ".join(current_path)
        rows.append(
            {
                "key": f"{name}::{_normalize_path(segment_path)}",
                "name": name,
                "path": segment_path,
                "description": str(pick(node, "description", "Description") or ""),
                "depth": max(0, len(current_path) - 1),
                "ranges": ranges,
                "child_count": sum(1 for child in children if isinstance(child, dict)),
            }
        )
    for child in children:
        if isinstance(child, dict):
            walk_admin_node(child, current_path, rows)


def admin_roots(payload: dict[str, Any] | list[Any]) -> list[Any]:
    if isinstance(payload, list):
        return payload
    if not isinstance(payload, dict):
        return []
    node = payload.get("node") or payload.get("Node")
    if isinstance(node, dict):
        return [node]
    if payload.get("name") or payload.get("Name"):
        return [payload]
    for key in ("segments", "Segments", "items", "Items", "nodes", "Nodes"):
        if isinstance(payload.get(key), list):
            return payload[key]
    return []


def find_admin_node(payload: dict[str, Any] | list[Any], segment_key: str, path: str, name: str) -> dict[str, Any] | None:
    target_key = str(segment_key or "").strip()
    target_path = _normalize_path(path)
    target_name = str(name or "").strip()

    def walk(node: dict[str, Any], parts: list[str], root_node: bool = False) -> dict[str, Any] | None:
        node_name = str(pick(node, "name", "Name") or "Unnamed segment").strip()
        display = "Segments" if root_node and node_name == "*root" else node_name
        current_parts = [*parts, display] if display else parts
        normalized_path = _normalize_path(" / ".join(current_parts))
        node_key = f"{node_name}::{normalized_path}"
        if not root_node and (
            (target_key and node_key == target_key)
            or (target_path and normalized_path == target_path and (not target_name or target_name == node_name))
            or (target_name and not target_path and target_name == node_name)
        ):
            return {"node": node, "key": node_key, "name": node_name, "path": normalized_path}
        for child in first_list(node, "nodes", "Nodes", "children", "Children"):
            if isinstance(child, dict):
                found = walk(child, current_parts)
                if found:
                    return found
        return None

    for root in admin_roots(payload):
        if isinstance(root, dict):
            found = walk(root, [], root_node=True)
            if found:
                return found
    return None


def normalize_hosts(payload: dict[str, Any]) -> list[dict[str, str]]:
    rows = first_list(payload, "hosts", "Hosts", "Host")
    if not rows and isinstance(payload.get("_embedded"), dict):
        for value in payload["_embedded"].values():
            if isinstance(value, list):
                rows = value
                break
    output = []
    seen = set()
    for row in rows:
        if not isinstance(row, dict):
            continue
        ip = str(pick(row, "ip", "IPAddress", "ipAddress", "ipv4") or "").strip()
        if not ip:
            ip = first_ipv4_value(row)
        if not ip:
            continue
        try:
            ip = str(ipaddress.IPv4Address(ip))
        except Exception:
            continue
        if ip in seen:
            continue
        seen.add(ip)
        output.append({"ip": ip, "id": str(pick(row, "hostId", "id", "ID") or ""), "mac": str(pick(row, "mac", "MACAddress", "macAddress") or "")})
    return output


def parse_host_ip_file(content: bytes) -> list[dict[str, str]]:
    text = content.decode("utf-8", errors="replace")
    try:
        payload = json.loads(text)
        if isinstance(payload, dict):
            if isinstance(payload.get("hosts"), list):
                return normalize_hosts({"hosts": payload["hosts"]})
            return normalize_hosts(payload)
        if isinstance(payload, list):
            return normalize_hosts({"hosts": payload})
    except json.JSONDecodeError:
        pass

    rows: list[dict[str, str]] = []
    seen: set[str] = set()

    def add_ip(value: str, row_id: str = "", mac: str = "") -> None:
        try:
            ip = str(ipaddress.IPv4Address(str(value).strip()))
        except Exception:
            return
        if ip in seen:
            return
        seen.add(ip)
        rows.append({"ip": ip, "id": row_id, "mac": mac})

    try:
        from io import StringIO

        reader = csv.DictReader(StringIO(text))
        if reader.fieldnames:
            for row in reader:
                candidates = [row.get(key, "") for key in ("ip", "IPAddress", "ipAddress", "ipv4", "IPv4")]
                candidates.extend(str(value) for value in row.values())
                for candidate in candidates:
                    before = len(rows)
                    add_ip(candidate, str(row.get("id", row.get("hostId", ""))), str(row.get("mac", "")))
                    if len(rows) > before:
                        break
    except Exception:
        pass

    if rows:
        return rows

    for match in re.findall(r"(?<!\d)(?:\d{1,3}\.){3}\d{1,3}(?!\d)", text):
        add_ip(match)
    if not rows:
        raise HTTPException(status_code=400, detail="No valid IPv4 host IPs were found in the uploaded host snapshot.")
    return rows


def offline_host_collector_script() -> str:
    return """#!/usr/bin/env python3
import argparse
import getpass
import ipaddress
import json
import ssl
import sys
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime, timezone


def open_url(request, verify_tls):
    context = None if verify_tls else ssl._create_unverified_context()
    with urllib.request.urlopen(request, timeout=45, context=context) as response:
        return response.read()


def login(base_url, username, password, verify_tls):
    body = urllib.parse.urlencode({\"username\": username, \"password\": password}).encode()
    request = urllib.request.Request(
        f\"{base_url.rstrip('/')}/api/login\",
        data=body,
        headers={\"Content-Type\": \"application/x-www-form-urlencoded\"},
        method=\"POST\",
    )
    token = open_url(request, verify_tls).decode(\"utf-8\", errors=\"replace\").strip()
    if not token:
        raise RuntimeError(\"Web API login returned an empty token\")
    return token


def find_ipv4(value):
    if isinstance(value, dict):
        for item in value.values():
            found = find_ipv4(item)
            if found:
                return found
    if isinstance(value, list):
        for item in value:
            found = find_ipv4(item)
            if found:
                return found
    if isinstance(value, str):
        try:
            return str(ipaddress.IPv4Address(value.strip()))
        except Exception:
            return \"\"
    return \"\"


def normalize_hosts(payload):
    rows = payload.get(\"hosts\") or payload.get(\"Hosts\") or payload.get(\"Host\") or []
    if not rows and isinstance(payload.get(\"_embedded\"), dict):
        for value in payload[\"_embedded\"].values():
            if isinstance(value, list):
                rows = value
                break
    output = []
    seen = set()
    for row in rows:
        if not isinstance(row, dict):
            continue
        ip = row.get(\"ip\") or row.get(\"IPAddress\") or row.get(\"ipAddress\") or row.get(\"ipv4\") or find_ipv4(row)
        if not ip:
            continue
        try:
            ip = str(ipaddress.IPv4Address(str(ip).strip()))
        except Exception:
            continue
        if ip in seen:
            continue
        seen.add(ip)
        output.append({\"ip\": ip, \"id\": str(row.get(\"hostId\") or row.get(\"id\") or \"\"), \"mac\": str(row.get(\"mac\") or row.get(\"MACAddress\") or \"\")})
    return output


def main():
    parser = argparse.ArgumentParser(description=\"Collect active host IPs for Segment Conflict Resolution Management.\")
    parser.add_argument(\"--url\", required=True, help=\"admin platform base URL, for example https://platform.example.local\")
    parser.add_argument(\"--username\", required=True)
    parser.add_argument(\"--password\", help=\"Optional. If omitted, the script prompts securely.\")
    parser.add_argument(\"--output\", default=\"host-ips.json\")
    parser.add_argument(\"--verify-tls\", action=\"store_true\", help=\"Verify the admin platform TLS certificate.\")
    args = parser.parse_args()

    password = args.password or getpass.getpass(\"Password: \")
    token = login(args.url, args.username, password, args.verify_tls)
    request = urllib.request.Request(
        f\"{args.url.rstrip('/')}/api/hosts\",
        headers={\"Authorization\": token, \"Accept\": \"application/hal+json, application/json\"},
        method=\"GET\",
    )
    payload = json.loads(open_url(request, args.verify_tls).decode(\"utf-8\"))
    hosts = normalize_hosts(payload)
    snapshot = {\"collected_at\": datetime.now(timezone.utc).isoformat(), \"count\": len(hosts), \"hosts\": hosts}
    with open(args.output, \"w\", encoding=\"utf-8\") as handle:
        json.dump(snapshot, handle, indent=2)
    print(f\"Wrote {len(hosts)} host IPs to {args.output}\")


if __name__ == \"__main__\":
    try:
        main()
    except Exception as exc:
        print(f\"ERROR: {exc}\", file=sys.stderr)
        raise SystemExit(1)
"""


def first_ipv4_value(value: Any) -> str:
    if isinstance(value, dict):
        for item in value.values():
            found = first_ipv4_value(item)
            if found:
                return found
    if isinstance(value, list):
        for item in value:
            found = first_ipv4_value(item)
            if found:
                return found
    if isinstance(value, str):
        try:
            return str(ipaddress.IPv4Address(value.strip()))
        except Exception:
            return ""
    return ""


def subtract_updates_from_ranges(ranges: list[str], updates: list[RangeUpdate]) -> list[str]:
    intervals = []
    for update in updates:
        parsed = parse_ip_range(update.remove_range)
        if parsed:
            intervals.append((parsed["start"], parsed["end"]))
    if not intervals:
        return ranges
    output = []
    for range_value in ranges:
        parsed_range = parse_ip_range(range_value)
        if not parsed_range:
            output.append(range_value)
            continue
        pieces = subtract_intervals(parsed_range["start"], parsed_range["end"], intervals)
        output.extend(format_ip_range(start, end) for start, end in pieces)
    return list(dict.fromkeys(output))


def subtract_intervals(start: int, end: int, intervals: list[tuple[int, int]]) -> list[tuple[int, int]]:
    removals = sorted((max(start, left), min(end, right)) for left, right in intervals if not (right < start or left > end))
    if not removals:
        return [(start, end)]
    merged = []
    for left, right in removals:
        if not merged or left > merged[-1][1] + 1:
            merged.append([left, right])
        else:
            merged[-1][1] = max(merged[-1][1], right)
    pieces = []
    cursor = start
    for left, right in merged:
        if cursor < left:
            pieces.append((cursor, left - 1))
        cursor = max(cursor, right + 1)
    if cursor <= end:
        pieces.append((cursor, end))
    return pieces


def parse_ip_range(value: str) -> dict[str, int] | None:
    raw = str(value or "").strip()
    try:
        if "/" in raw:
            network = ipaddress.ip_network(raw, strict=False)
            if network.version != 4:
                return None
            return {"start": int(network.network_address), "end": int(network.broadcast_address)}
        if "-" in raw:
            left, right = [part.strip() for part in raw.split("-", 1)]
            start, end = int(ipaddress.IPv4Address(left)), int(ipaddress.IPv4Address(right))
        else:
            start = end = int(ipaddress.IPv4Address(raw))
        if start > end:
            start, end = end, start
        return {"start": start, "end": end}
    except Exception:
        return None


def format_ip_range(start: int, end: int) -> str:
    left = str(ipaddress.IPv4Address(start))
    right = str(ipaddress.IPv4Address(end))
    return left if start == end else f"{left}-{right}"


def _ip_sort_token(value: str) -> tuple[int, int]:
    parsed = parse_ip_range(value)
    if not parsed:
        return (0, 0)
    return (int(parsed["start"]), int(parsed["end"]))


def authoritative_segment(left: dict[str, Any], right: dict[str, Any]) -> dict[str, Any]:
    candidates = [left, right]
    candidates.sort(key=lambda row: (-(row.get("depth", 0) or 0), len(row.get("ranges") or []), row.get("path", ""), row.get("name", "")))
    return candidates[0]


def rule_segment_refs(node: ET.Element, source: str) -> list[dict[str, str]]:
    refs = []
    for segment in direct_children(node, "SEGMENT"):
        refs.append({"id": segment.attrib.get("ID", ""), "name": segment.attrib.get("NAME", ""), "source": source})
    return refs


def unique_refs(refs: list[dict[str, str]]) -> list[dict[str, str]]:
    output = []
    seen = set()
    for ref in refs:
        key = (ref.get("id", ""), ref.get("name", ""), ref.get("source", ""))
        if key in seen:
            continue
        seen.add(key)
        output.append(ref)
    return output


def parse_xml(content: bytes) -> ET.Element:
    try:
        return ET.fromstring(content)
    except ET.ParseError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid XML: {exc}") from exc


def local_name(tag: Any) -> str:
    return str(tag).rsplit("}", 1)[-1]


def direct_child(node: ET.Element | None, name: str) -> ET.Element | None:
    if node is None:
        return None
    for child in list(node):
        if local_name(child.tag) == name:
            return child
    return None


def direct_children(node: ET.Element | None, name: str) -> list[ET.Element]:
    if node is None:
        return []
    return [child for child in list(node) if local_name(child.tag) == name]


def iter_tag(node: ET.Element | None, name: str):
    if node is None:
        return
    for child in node.iter():
        if local_name(child.tag) == name:
            yield child


def bool_attr(value: str | None, default: bool) -> bool:
    return default if value is None else value.lower() == "true"


def load_parsed_artifact(kind: str) -> dict[str, Any]:
    path = UPLOAD_DIR / f"{kind}.json"
    return read_json(path, {"summary": {}, kind: []})


def read_config() -> dict[str, Any]:
    return read_json(CONFIG_PATH, {})


def write_config(config: dict[str, Any]) -> None:
    CONFIG_PATH.write_text(json.dumps(config, indent=2), encoding="utf-8")


def saved_connection(kind: str) -> ConnectionConfig:
    saved = read_config().get(kind, {})
    password = reveal_password(saved.get("password", ""))
    if not saved.get("base_url") or not saved.get("username") or not password:
        raise HTTPException(status_code=400, detail=f"{kind} API configuration is incomplete.")
    return ConnectionConfig(
        base_url=saved["base_url"],
        username=saved["username"],
        password=password,
        verify_tls=bool(saved.get("verify_tls", False)),
    )


def connection_from_payload_or_saved(kind: str, payload: ApiConfigPayload | None) -> ConnectionConfig:
    if payload and (payload.base_url or payload.username or payload.password):
        saved = read_config().get(kind, {})
        return ConnectionConfig(
            base_url=payload.base_url or saved.get("base_url", ""),
            username=payload.username or saved.get("username", ""),
            password=payload.password or reveal_password(saved.get("password", "")),
            verify_tls=bool(payload.verify_tls),
        )
    return saved_connection(kind)


def masked_config(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "base_url": row.get("base_url", ""),
        "username": row.get("username", ""),
        "password_saved": bool(reveal_password(row.get("password", ""))),
        "verify_tls": bool(row.get("verify_tls", False)),
        "updated_at": row.get("updated_at", ""),
    }


def hide_password(value: str) -> str:
    if not value:
        return ""
    return base64.b64encode(value.encode("utf-8")).decode("ascii")


def reveal_password(value: str) -> str:
    if not value:
        return ""
    try:
        return base64.b64decode(value.encode("ascii")).decode("utf-8")
    except Exception:
        return ""


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def file_meta(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"exists": False}
    stat = path.stat()
    return {"exists": True, "name": path.name, "size": stat.st_size, "updated_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat()}


def _open_url(request: urllib.request.Request, label: str, api_name: str, timeout: int, context: ssl.SSLContext | None) -> bytes:
    last_error: Exception | None = None
    for attempt in range(1, 3):
        try:
            with urllib.request.urlopen(request, timeout=timeout, context=context) as response:
                return response.read()
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            raise PlatformApiError(f"{api_name} {label} failed with HTTP {exc.code}: {detail or exc.reason}") from exc
        except urllib.error.URLError as exc:
            last_error = exc
            if attempt == 2 or not is_transient(exc.reason):
                raise PlatformApiError(f"Unable to reach {api_name}: {exc.reason}") from exc
        except (TimeoutError, socket.timeout) as exc:
            last_error = exc
            if attempt == 2:
                raise PlatformApiError(f"{api_name} request timed out.") from exc
        time.sleep(0.5 * attempt)
    raise PlatformApiError(f"{api_name} request failed: {last_error}")


def _token_from_response(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace").strip()
    if not text:
        return ""
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return text
    if isinstance(payload, dict):
        for key in ("access_token", "token", "id_token"):
            if payload.get(key):
                return str(payload[key]).strip()
    if isinstance(payload, str):
        return payload.strip()
    return ""


def is_transient(reason: Any) -> bool:
    return isinstance(reason, (TimeoutError, socket.timeout)) or "timed out" in str(reason).lower()


def first_list(payload: dict[str, Any], *keys: str) -> list[Any]:
    for key in keys:
        if isinstance(payload.get(key), list):
            return payload[key]
    return []


def pick(row: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        if key in row:
            return row.get(key)
    return None


def _normalize_path(path: str) -> str:
    parts = [part.strip().casefold() for part in str(path or "").replace(">", "/").split("/") if part.strip()]
    if parts and parts[0] == "segments":
        parts = parts[1:]
    return "/".join(parts)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()
